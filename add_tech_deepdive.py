# pyre-ignore-all-errors
"""
Add Technology Deep Dive section to SYSTEM_ARCHITECTURE.docx
Each tech: Why Selected + Black-Box + Internal Working
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc_path = r"c:\Users\L Melvin Denish\barclays\PreDelinquencyEngine\SYSTEM_ARCHITECTURE.docx"
doc = Document(doc_path)

doc.add_page_break()

# ════════════════════════════════════════
# MAIN HEADING
# ════════════════════════════════════════
h = doc.add_heading("23. Technology Deep Dive", level=1)
doc.add_paragraph(
    "This section explains every technology in the stack: why it was chosen over alternatives, "
    "how it works at a glance (black-box), and how it operates internally."
)

# ─── Helper ───
def add_tech(name, why, blackbox, internal, alternatives=""):
    doc.add_heading(name, level=2)
    
    doc.add_heading("Why Selected", level=3)
    doc.add_paragraph(why)
    
    if alternatives:
        doc.add_heading("Alternatives Considered", level=3)
        doc.add_paragraph(alternatives)
    
    doc.add_heading("Black-Box (Quick Read)", level=3)
    doc.add_paragraph(blackbox)
    
    doc.add_heading("Internal Working", level=3)
    doc.add_paragraph(internal)

# ════════════════════════════════════════
# EACH TECHNOLOGY
# ════════════════════════════════════════

add_tech(
    "23.1 Apache Kafka",
    
    "Kafka was selected as our message broker because the Pre-Delinquency Engine requires "
    "real-time transaction streaming with guaranteed delivery and ordering. In banking, "
    "losing even one transaction event can mean missing a critical stress signal. Kafka provides:\n"
    "• Exactly-once delivery semantics — no duplicate or lost transactions\n"
    "• Topic partitioning by customer_id — ensures all transactions for one customer go to the same partition, maintaining order\n"
    "• High throughput — handles 100K+ messages/second, essential when processing transactions from 1000+ customers\n"
    "• Persistent log — messages are retained for 7 days, allowing replay for debugging or reprocessing\n"
    "• Decoupling — data generators, Flink, and Debezium all communicate through Kafka without knowing about each other",
    
    "Alternatives: RabbitMQ (good for task queues but lacks persistent log and replay), "
    "AWS SQS (vendor lock-in), Redis Pub/Sub (no persistence, messages lost if consumer is down). "
    "Kafka's persistent log + consumer group model is ideal for event sourcing in banking.",
    
    "Think of Kafka as a postal system:\n"
    "• Producers (transaction generator) drop letters (messages) into mailboxes (topics)\n"
    "• Each mailbox has sections (partitions) — letters with the same customer_id go to the same section\n"
    "• Consumers (Flink, Debezium) pick up letters from their assigned sections\n"
    "• Letters stay in the mailbox for 7 days even after being read, so others can read them too\n\n"
    "Input: Transaction JSON → Output: Message delivered to all subscribed consumers",
    
    "Kafka is a distributed commit log. Internally:\n\n"
    "1. PRODUCERS: When transaction_generator publishes a transaction, it serializes the dict to JSON "
    "and sends it to Kafka with the customer_id as the key. Kafka hashes the key (murmur2 hash) "
    "to determine which partition the message goes to. The producer batches messages (linger.ms=5ms) "
    "and compresses them (snappy) before sending to reduce network overhead.\n\n"
    "2. BROKER STORAGE: Kafka stores messages as an append-only log on disk. Each partition is a "
    "sequence of segment files (1GB each). Messages have an offset (sequential ID). The broker "
    "uses page cache (OS memory) for reads, making sequential reads nearly as fast as memory. "
    "This is why Kafka is fast despite writing to disk.\n\n"
    "3. CONSUMERS: Flink's Kafka consumer maintains a consumer offset (how far it has read). "
    "If Flink crashes and restarts, it resumes from the last committed offset — no data loss. "
    "Multiple consumer groups can read the same topic independently. Our setup has: "
    "group 'flink-pdi' for stream processing and group 'debezium' for CDC.\n\n"
    "4. TOPICS IN OUR ENGINE:\n"
    "   • pdi-transactions: Raw transaction events (from generator + Debezium)\n"
    "   • pdi-account-updates: Daily balance snapshots\n"
    "   • pdi.public.customers: CDC events from customers table\n"
    "   • pdi.public.risk_scores: CDC events from risk_scores table\n\n"
    "5. CONFIGURATION: We use Confluent Platform (cp-kafka:7.6.0) with:\n"
    "   • KAFKA_OFFSETS_TOPIC_REPLICATION_FACTOR=1 (single broker for dev)\n"
    "   • KAFKA_LISTENER_SECURITY_PROTOCOL_MAP separating internal (Docker) and external (host) traffic\n"
    "   • Port 29092 exposed for host applications, port 9092 for internal Docker communication"
)

add_tech(
    "23.2 Apache Flink",
    
    "Flink was chosen for real-time stream processing because it provides:\n"
    "• True streaming (not micro-batching like Spark Streaming) — processes events one-by-one with < 10ms latency\n"
    "• Stateful processing — maintains per-customer rolling windows (7d, 30d) in memory with checkpointing\n"
    "• Event time processing — uses transaction timestamps, not arrival time, ensuring correct window computations even if events arrive late\n"
    "• Exactly-once semantics — integrated with Kafka for end-to-end exactly-once processing\n"
    "• Fault tolerance via checkpoints — if a task fails, it restores state from the last checkpoint and replays from Kafka",
    
    "Alternatives: Spark Structured Streaming (micro-batch, higher latency ~500ms), "
    "Kafka Streams (simpler but can't handle complex windowed aggregations as cleanly), "
    "AWS Kinesis (vendor lock-in). Flink's true streaming + stateful windowing is ideal "
    "for computing rolling behavioral features.",
    
    "Think of Flink as a conveyor belt in a factory:\n"
    "• Transactions arrive on the belt one by one from Kafka\n"
    "• Each transaction passes through processing stations (operators)\n"
    "• Station 1: Groups by customer_id (keyBy)\n"
    "• Station 2: Updates the 7-day and 30-day counters for that customer\n"
    "• Station 3: Writes updated features to PostgreSQL\n\n"
    "Input: Raw transaction events → Output: Updated streaming_features per customer",
    
    "Flink's internal architecture:\n\n"
    "1. JOB GRAPH: Our Flink job (stream_processing/flink_job.py) is compiled into a directed acyclic "
    "graph (DAG) of operators:\n"
    "   Source(Kafka) → KeyBy(customer_id) → ProcessFunction(compute_features) → Sink(PostgreSQL)\n\n"
    "2. TASK SLOTS: The JobManager distributes the DAG to TaskManagers. Each TaskManager has slots "
    "(threads). Our setup: 1 TaskManager with 4 slots, handling up to 4 parallel keys.\n\n"
    "3. STATE MANAGEMENT: For each customer_id, Flink maintains state:\n"
    "   • HashMap<customer_id, {txn_list_7d, txn_list_30d, last_balance, counters}>\n"
    "   • State is stored in RocksDB (embedded key-value store) and periodically checkpointed "
    "to persistent storage. If the job crashes, it restores from the last checkpoint.\n\n"
    "4. WINDOW COMPUTATIONS: When a new transaction arrives for customer_id='CUST_ABC':\n"
    "   a. Flink routes it to the correct slot via hash(customer_id)\n"
    "   b. The ProcessFunction adds it to the 7d and 30d lists\n"
    "   c. It evicts transactions older than 7/30 days from the lists\n"
    "   d. It recomputes: discretionary_spend_7d = sum(amount for txn in 7d_list if category in disc_cats)\n"
    "   e. Similarly for atm_count, lending_app_count, failed_autodebits, etc.\n"
    "   f. Upserts the result to streaming_features table\n\n"
    "5. LOCAL FALLBACK: If Flink cluster is unavailable, our system falls back to a local Python "
    "Kafka consumer (stream_processing/flink_job.py::create_flink_job_local) that mimics the same "
    "logic using in-memory dictionaries. This is what runs in the demo."
)

add_tech(
    "23.3 Apache Spark",
    
    "Spark was selected for batch feature computation because:\n"
    "• Distributed processing — can parallelize across 1000+ customers simultaneously\n"
    "• DataFrame API — SQL-like operations on large transaction datasets\n"
    "• Integration with PostgreSQL — reads directly via JDBC\n"
    "• Handles complex aggregations — monthly spend averages, volatility calculations, "
    "trend analysis across 6-month windows that are too expensive for real-time processing",
    
    "Alternatives: Plain Pandas (works for 1000 customers but won't scale to millions), "
    "Dask (good but less mature ecosystem), BigQuery (cloud-only). "
    "Spark provides the best balance of scalability and local development support.",
    
    "Spark as a calculator with multiple screens:\n"
    "• You give it a big spreadsheet (transactions table from PostgreSQL)\n"
    "• It splits the spreadsheet across multiple screens (workers)\n"
    "• Each screen calculates salary delays, spend trends for its chunk of customers\n"
    "• Results are merged and written back to PostgreSQL batch_features table\n\n"
    "Input: 6 months of transactions → Output: Historical features per customer",
    
    "Spark's internal process for our batch job:\n\n"
    "1. DRIVER: The Spark Driver (batch_processing/spark_jobs.py) creates a SparkSession and "
    "reads the customers and transactions tables from PostgreSQL via JDBC.\n\n"
    "2. LAZY EVALUATION: Spark builds a logical plan (DAG) of transformations but doesn't "
    "execute until an action (like .toPandas() or .write) is called. This allows Spark to "
    "optimize the execution order.\n\n"
    "3. PARTITIONING: The transactions DataFrame is partitioned by customer_id. Each partition "
    "is processed by a separate executor (worker). With 1000 customers and 4 cores, each core "
    "handles ~250 customers.\n\n"
    "4. FEATURE COMPUTATION per partition:\n"
    "   a. salary_delay_days: For each customer, find salary_credit transactions in last 90 days, "
    "compare actual day vs expected salary_credit_day from customers table\n"
    "   b. utility_payment_delay_avg: Similar for utility bills vs their expected due date\n"
    "   c. discretionary_spend_trend: ratio of last 7 days discretionary spend vs previous 30-day average\n"
    "   d. avg_monthly_spend_3m: Group debits by month, compute average\n"
    "   e. spend_volatility_3m: std(monthly_spend) / mean(monthly_spend) — high volatility = instability\n\n"
    "5. LOCAL FALLBACK: When Spark cluster is unavailable (demo mode), main.py::_compute_features_locally() "
    "runs the identical logic using Pandas DataFrames. Same features, same calculations, just single-threaded."
)

add_tech(
    "23.4 PostgreSQL",
    
    "PostgreSQL was chosen as our primary database because:\n"
    "• ACID compliance — critical for banking data (no partial writes, no data corruption)\n"
    "• JSONB support — stores SHAP drivers and product_holdings as JSON efficiently\n"
    "• Array types — TEXT[] for product_holdings allows multi-value fields\n"
    "• Logical replication — required for Debezium CDC (WAL-based change capture)\n"
    "• Rich indexing — B-tree, GiST, GIN indexes for fast lookups on customer_id, timestamp\n"
    "• Mature ecosystem — psycopg2 driver, SQLAlchemy ORM, proven in banking for decades",
    
    "Alternatives: MySQL (less advanced JSON/array support), MongoDB (no ACID for multi-document), "
    "CockroachDB (overkill for single-node), SQLite (no concurrent writes for multi-service). "
    "PostgreSQL's WAL-based replication + rich type system makes it the clear choice.",
    
    "PostgreSQL as a filing cabinet:\n"
    "• 12 drawers (tables): customers, transactions, risk_scores, notifications, etc.\n"
    "• Each drawer has organized folders (indexes) for quick lookup\n"
    "• When you add/change a file, a log entry is written (WAL) that Debezium watches\n\n"
    "Input: SQL queries → Output: Structured data rows",
    
    "PostgreSQL's role in our architecture:\n\n"
    "1. WRITE-AHEAD LOG (WAL): Every INSERT/UPDATE is first written to the WAL (a sequential "
    "log file on disk) before the actual table. This guarantees durability — if the server "
    "crashes after WAL write but before table write, PostgreSQL replays the WAL on restart. "
    "Debezium reads this WAL for CDC (Change Data Capture).\n\n"
    "2. CONNECTION POOLING: Multiple services connect simultaneously — data generator writes "
    "customers, scoring service reads features, notification dispatcher writes to rm_tasks. "
    "PostgreSQL handles this via multi-process architecture (one backend process per connection).\n\n"
    "3. INDEXING: We create B-tree indexes on:\n"
    "   • transactions(customer_id) — 37K rows, O(log n) lookup instead of O(n)\n"
    "   • transactions(timestamp) — for time-range queries in batch processing\n"
    "   • risk_scores(customer_id, scored_at) — latest score lookup\n"
    "   • notifications(customer_id, channel, status) — cooldown checks\n\n"
    "4. LOGICAL REPLICATION: PostgreSQL is configured with wal_level=logical and creates "
    "a replication slot 'pdi_slot' for Debezium. This means Debezium receives a stream "
    "of ROW-level changes (INSERT, UPDATE, DELETE) without polling.\n\n"
    "5. OUR SCHEMA: 12 tables (see Section 4.3) with foreign key relationships. "
    "customers is the central table; transactions, risk_scores, notifications, "
    "rm_tasks, collector_assignments all reference it via customer_id."
)

add_tech(
    "23.5 XGBoost (Extreme Gradient Boosting)",
    
    "XGBoost is our primary ML model because:\n"
    "• Best-in-class for tabular data — consistently wins Kaggle competitions on structured data\n"
    "• Built-in handling of missing values — essential when some features are unavailable\n"
    "• Regularization (L1 + L2) — prevents overfitting on 1000-customer dataset\n"
    "• SHAP integration — TreeSHAP provides EXACT Shapley values in O(TLD) time (fast)\n"
    "• Interpretable feature importance — critical for banking (regulators require explainability)\n"
    "• Handles class imbalance — scale_pos_weight parameter balances delinquent vs non-delinquent",
    
    "Alternatives: Random Forest (less accurate, no gradient boosting), "
    "CatBoost (good but less SHAP support), Neural Networks (black-box, hard to explain to regulators), "
    "Logistic Regression (too simple for complex feature interactions). "
    "XGBoost + SHAP gives us both accuracy AND explainability — critical for FCA/PRA regulatory compliance.",
    
    "XGBoost as a committee of experts:\n"
    "• 300 experts (decision trees) are trained sequentially\n"
    "• Each new expert focuses on what the previous experts got WRONG (gradient boosting)\n"
    "• To predict, all 300 experts vote, and votes are weighted and summed\n"
    "• Output: probability between 0 and 1 (risk of delinquency)\n\n"
    "Input: 40 feature values → Output: 0.0–1.0 risk score",
    
    "XGBoost's internal mechanics:\n\n"
    "1. GRADIENT BOOSTING: At each iteration t, XGBoost fits a new tree to the RESIDUALS "
    "(errors) of the previous ensemble. Mathematically:\n"
    "   F_t(x) = F_{t-1}(x) + η × h_t(x)\n"
    "   where η=0.05 (learning rate) and h_t is the new tree fit to negative gradients.\n\n"
    "2. TREE CONSTRUCTION: Each tree is built top-down by finding the best split at each node. "
    "For each feature, XGBoost computes the gain:\n"
    "   Gain = 0.5 × [G_L²/H_L + G_R²/H_R - (G_L+G_R)²/(H_L+H_R)] - γ\n"
    "   where G=sum of gradients, H=sum of hessians, γ=regularization penalty.\n"
    "   This prefers splits that maximally separate delinquent from non-delinquent customers.\n\n"
    "3. REGULARIZATION: Our config uses:\n"
    "   • max_depth=6: trees can't become too complex\n"
    "   • reg_alpha=0.1 (L1): encourages sparse trees (some features get zero importance)\n"
    "   • reg_lambda=1.0 (L2): shrinks leaf weights toward zero\n"
    "   • subsample=0.8: each tree sees only 80% of data (reduces overfitting)\n"
    "   • colsample_bytree=0.8: each tree uses only 80% of features\n\n"
    "4. PREDICTION: For a customer with features x, XGBoost sums all 300 trees:\n"
    "   raw_score = sum(tree_t.predict(x) for t in 1..300)\n"
    "   probability = sigmoid(raw_score) = 1 / (1 + exp(-raw_score))\n\n"
    "5. OUR RESULTS: AUC=0.896 with top features being:\n"
    "   failed_autodebits_count_30d, lending_app_txn_count_7d, salary_delay_days, "
    "   discretionary_spend_trend, savings_balance_pct_change_7d"
)

add_tech(
    "23.6 LightGBM (Light Gradient Boosting Machine)",
    
    "LightGBM serves as our validation model and second opinion:\n"
    "• 10× faster training than XGBoost (histogram-based splitting)\n"
    "• Leaf-wise growth (vs XGBoost's level-wise) — deeper trees with fewer nodes\n"
    "• Native categorical feature support — handles income_bracket, region without encoding\n"
    "• GOSS (Gradient-based One-Side Sampling) — keeps all high-gradient samples, subsamples low-gradient ones\n"
    "• Cross-validates XGBoost's findings — if both agree on feature importance, we're confident",
    
    "Selected alongside XGBoost (not instead of) because having two different tree algorithms "
    "agreeing on the same risk drivers provides robustness. If XGBoost says 'salary_delay is important' "
    "and LightGBM agrees, we have high confidence in that signal.",
    
    "Same committee analogy as XGBoost, but with a different strategy:\n"
    "• Instead of growing trees level-by-level evenly, LightGBM grows the leaf with the highest error first\n"
    "• This produces deeper, more specialized trees with fewer total splits\n"
    "• Faster training because it uses histograms (bins) instead of sorting every feature value\n\n"
    "Input: Same 40 features → Output: 0.0–1.0 risk score (should agree closely with XGBoost)",
    
    "LightGBM's key internal differences from XGBoost:\n\n"
    "1. HISTOGRAM-BASED SPLITTING: Instead of evaluating every unique feature value as a split "
    "candidate (expensive), LightGBM bins continuous features into 255 buckets and evaluates "
    "only bucket boundaries. This reduces split evaluation from O(n×features) to O(255×features).\n\n"
    "2. LEAF-WISE GROWTH: XGBoost grows trees level-by-level (all nodes at depth d before depth d+1). "
    "LightGBM grows the leaf with the highest loss reduction first. This creates asymmetric trees "
    "that are deeper on the side that matters most.\n\n"
    "3. GOSS: During training, LightGBM keeps ALL samples with large gradients (high error — "
    "hard cases) and randomly samples from small gradient samples (easy cases). This speeds "
    "training while preserving accuracy on the difficult cases.\n\n"
    "4. IN OUR ENGINE: LightGBM achieves AUC=0.896 (matching XGBoost), confirming that the "
    "same features drive risk. Its feature importance ranking is used to cross-validate SHAP results."
)

add_tech(
    "23.7 PyTorch + LSTM (Long Short-Term Memory)",
    
    "We chose PyTorch for the LSTM because:\n"
    "• Dynamic computation graphs — easier to debug than TensorFlow's static graphs\n"
    "• Native Python feel — tensors work like numpy arrays\n"
    "• LSTM is ideal for sequence data — captures how customer behavior CHANGES over time\n"
    "• Complementary to tree models — captures temporal patterns that XGBoost/LightGBM miss\n"
    "• PyTorch is the standard for research and increasingly for production",
    
    "Alternatives: TensorFlow/Keras (larger footprint, harder debugging), "
    "Transformer models (overkill for 30-timestep sequences), "
    "1D-CNN (faster but less effective at capturing long-range dependencies). "
    "LSTM + PyTorch gives the best balance of temporal modeling capability and code simplicity.",
    
    "LSTM as a person reading a customer's transaction diary:\n"
    "• They read 30 days of transactions, one day at a time\n"
    "• Each day, they update their 'memory' of the customer's behavior\n"
    "• They selectively REMEMBER important events (EMI bounce) and FORGET noise (normal grocery)\n"
    "• After reading all 30 days, they make a judgment: 'this customer is trending toward risk'\n\n"
    "Input: 30 days × 40 features per day → Output: 0.0–1.0 probability",
    
    "LSTM internals in our model:\n\n"
    "1. ARCHITECTURE: DelinquencyLSTM(\n"
    "   LSTM(input=40, hidden=64, layers=2, dropout=0.3)\n"
    "   → Linear(64→32) → ReLU → Dropout(0.2) → Linear(32→1) → Sigmoid\n"
    ")\n\n"
    "2. LSTM CELL MECHANICS: At each timestep t (each day), the LSTM cell computes:\n"
    "   • Forget gate: f_t = σ(W_f · [h_{t-1}, x_t] + b_f)\n"
    "     Decides what to forget from previous memory (e.g., a normal grocery transaction)\n"
    "   • Input gate: i_t = σ(W_i · [h_{t-1}, x_t] + b_i)\n"
    "     Decides what new info to store (e.g., an EMI bounce is very memorable)\n"
    "   • Cell state update: C_t = f_t × C_{t-1} + i_t × tanh(W_c · [h_{t-1}, x_t] + b_c)\n"
    "     Old memory × forget + new memory × input\n"
    "   • Output gate: o_t = σ(W_o · [h_{t-1}, x_t] + b_o)\n"
    "   • Hidden state: h_t = o_t × tanh(C_t)\n"
    "     This h_t is passed to the next timestep AND to the output layer at the final step.\n\n"
    "3. TRAINING: BCEWithLogitsLoss (binary cross-entropy) with Adam optimizer (lr=0.001). "
    "30 epochs, batch_size=64. Gradient clipping prevents exploding gradients.\n\n"
    "4. WHY IT HELPS: Consider a customer whose spending gradually increases over 3 weeks before "
    "an EMI bounce. XGBoost sees a snapshot (high spend + bounce). LSTM sees the TRAJECTORY "
    "(spend creeping up → ATM spike → balance drop → bounce). The trajectory is a stronger "
    "predictive signal than any single snapshot.\n\n"
    "5. OUR RESULTS: LSTM achieves AUC=0.741 standalone (lower than XGBoost because it works "
    "better with more data). But in the ensemble, it contributes unique temporal signals that "
    "XGBoost misses, especially for customers with gradual deterioration patterns."
)

add_tech(
    "23.8 SHAP (SHapley Additive exPlanations)",
    
    "SHAP is mandatory in banking ML for regulatory compliance (FCA, PRA, GDPR Article 22):\n"
    "• Provides mathematically rigorous explanations for every individual prediction\n"
    "• TreeSHAP is exact (not approximate) for tree models — polynomial time\n"
    "• Drives our intervention routing — the TOP SHAP feature determines what TYPE of help to offer\n"
    "• Required for model audit trails — every score must be explainable",
    
    "Alternatives: LIME alone (approximate, inconsistent), "
    "Feature importance (only global, not per-prediction), "
    "Attention weights (only for neural nets). "
    "SHAP is the gold standard for ML explainability in regulated industries.",
    
    "SHAP as a 'blame allocation' system:\n"
    "• For each prediction, SHAP distributes the 'blame' across all 40 features\n"
    "• 'salary_delay_days contributed +0.18 to your risk score'\n"
    "• 'having_fixed_deposit contributed -0.05 (reduced your risk)'\n"
    "• All SHAP values sum to: prediction - average_prediction\n\n"
    "Input: Model + single customer features → Output: SHAP value per feature",
    
    "SHAP's mathematical foundation:\n\n"
    "1. SHAPLEY VALUES: From cooperative game theory. The contribution of feature i is the "
    "average marginal contribution across ALL possible orderings of features. For 40 features, "
    "that's 40! orderings — computationally intractable.\n\n"
    "2. TREESHAP: For tree models, Shapley values can be computed in polynomial time by "
    "exploiting the tree structure. It traverses each tree once, tracking how each feature's "
    "value routes the sample through branches. This gives EXACT values in O(TLD) time "
    "(T=trees, L=leaves, D=depth) instead of O(2^40).\n\n"
    "3. IN OUR ENGINE: After scoring a customer, we call:\n"
    "   shap_values = explainer.shap_values(customer_features)\n"
    "   top_drivers = sorted(zip(features, shap_values), key=abs, reverse=True)[:5]\n"
    "   This tells us: 'salary_delay_days pushed risk up by +0.183, lending_app_count pushed "
    "   up by +0.142, fixed_deposit pushed down by -0.051'\n\n"
    "4. INTERVENTION ROUTING: The rules engine maps top SHAP feature to intervention type:\n"
    "   salary_delay → payment_holiday_offer\n"
    "   lending_app_usage → wellness_checkin\n"
    "   failed_autodebits → emi_restructuring\n"
    "   This makes interventions contextually relevant, not generic."
)

add_tech(
    "23.9 FastAPI",
    
    "FastAPI was selected for the scoring REST API because:\n"
    "• Async support — handles concurrent scoring requests without blocking\n"
    "• Auto-generated Swagger/OpenAPI docs — instant API documentation at /docs\n"
    "• Pydantic validation — request/response schemas are auto-validated\n"
    "• Performance — one of the fastest Python web frameworks (built on Starlette + Uvicorn)\n"
    "• Type hints — full IDE autocompletion and error detection",
    
    "Alternatives: Flask (synchronous, slower), Django (too heavy for a microservice), "
    "gRPC (better for inter-service but we need REST for n8n/browser). "
    "FastAPI is the modern standard for Python ML serving.",
    
    "FastAPI as a drive-through window:\n"
    "• Customer sends request: 'Score customer CUST_ABC123'\n"
    "• FastAPI receives it, validates the format\n"
    "• Passes to the kitchen (ML models)\n"
    "• Returns the meal (risk score + SHAP explanations) in < 100ms\n\n"
    "Input: POST /score {customer_id} → Output: JSON with score, tier, SHAP, ensemble breakdown",
    
    "FastAPI scoring flow (12 internal steps):\n\n"
    "1. Uvicorn (ASGI server) receives HTTP POST on port 8000\n"
    "2. FastAPI validates request body against ScoreRequest Pydantic model\n"
    "3. Checks Redis cache: has this customer been scored in the last 5 minutes?\n"
    "4. If not cached, queries streaming_features table for real-time features\n"
    "5. Queries batch_features table for historical features\n"
    "6. Merges into a single numpy array of 40+ features\n"
    "7. XGBoost model.predict_proba(features) → xgb_score\n"
    "8. LightGBM model.predict_proba(features) → lgb_score\n"
    "9. LSTM model(sequence) → lstm_score (if temporal data available)\n"
    "10. Weighted ensemble: final = 0.6×xgb + 0.4×lstm\n"
    "11. SHAP TreeExplainer computes top 5 feature contributions\n"
    "12. Result stored in: risk_scores table + Redis cache + Cassandra\n"
    "13. JSON response returned in < 100ms\n\n"
    "The app also exposes /health (model status), /metrics (Prometheus), and /docs (Swagger UI)."
)

add_tech(
    "23.10 Redis",
    
    "Redis serves multiple roles in our architecture:\n"
    "• Score caching — avoids re-scoring the same customer within 5 minutes\n"
    "• Risk tier tracking — stores previous risk tier per customer for transition detection\n"
    "• Notification cooldown — prevents SMS/email spam (24-hour cooldown per channel)\n"
    "• Celery broker — task queue for asynchronous intervention processing\n"
    "• Sub-millisecond reads — in-memory data store, fastest possible lookups",
    
    "Alternatives: Memcached (no persistence, no pub/sub), DynamoDB (cloud-only, cost). "
    "Redis's in-memory speed + data structures (SET, SETEX for TTL) make it ideal for "
    "caching and state tracking.",
    
    "Redis as a sticky-note board:\n"
    "• Each sticky note has a key and value: 'risk_tier:CUST_ABC = critical'\n"
    "• Some notes have expiry timers (TTL) — score cache expires after 5 min\n"
    "• Reading a note is near-instant (microseconds)\n\n"
    "Input: GET/SET commands → Output: String/Hash values",
    
    "Redis in our engine:\n\n"
    "1. RISK TIER TRACKING (rules_engine.py):\n"
    "   Key: risk_tier:CUST_ABC → Value: 'watch'\n"
    "   When new score produces 'critical', rules engine detects transition: watch→critical → trigger alert\n\n"
    "2. NOTIFICATION COOLDOWN (notification_dispatcher.py):\n"
    "   Key: notif_cooldown:CUST_ABC:sms → Value: '1' → TTL: 86400 seconds (24 hours)\n"
    "   Before sending SMS, check: if key exists → skip (cooldown active)\n\n"
    "3. SCORE CACHE (scoring_service/app.py):\n"
    "   Key: score_cache:CUST_ABC → Value: JSON with score, tier, SHAP → TTL: 300 seconds\n"
    "   Prevents expensive re-computation for rapid successive requests.\n\n"
    "4. DATA STRUCTURE: Redis is single-threaded event loop (epoll-based). All operations "
    "are atomic. SETEX(key, ttl, value) sets a key with automatic expiration in one atomic operation."
)

add_tech(
    "23.11 Groq LLM (Llama 3.3 70B)",
    
    "Groq was selected for GenAI message generation because:\n"
    "• Fastest LLM inference — 300+ tokens/second on Llama 3.3 70B\n"
    "• Free tier — sufficient for hackathon demo (14,400 requests/day)\n"
    "• Llama 3.3 70B — open-weight model, no vendor lock-in concerns\n"
    "• Low latency — messages generated in < 1 second (vs 5-10s for OpenAI)\n"
    "• Supports system prompts — allows us to enforce Barclays brand voice",
    
    "Alternatives: OpenAI GPT-4 (expensive, 5-10s latency), Anthropic Claude (expensive), "
    "Local Ollama (requires GPU, slow on CPU). Groq's LPU architecture provides unmatched "
    "speed for real-time message generation during customer scoring.",
    
    "Groq as a writing assistant:\n"
    "• We give it: customer context + intervention type + channel\n"
    "• It returns: a personalized, empathetic message tailored to that channel\n"
    "• SMS → 160 chars, Email → 5 sentences, RM → full call script\n\n"
    "Input: System prompt + customer context → Output: Channel-specific message text",
    
    "Groq LLM in our notification pipeline:\n\n"
    "1. SYSTEM PROMPTS: Each channel has a carefully crafted system prompt:\n"
    "   • SMS: 'Write under 160 chars, warm, supportive, never accusatory, Indian English'\n"
    "   • Email: '3-5 sentences, professional, include offer details'\n"
    "   • RM Script: 'Structured brief with talking points, objection handling'\n"
    "   • Collector: 'Factual, structured, empathetic pre-delinquency tone'\n\n"
    "2. CONTEXT INJECTION: We build a prompt with:\n"
    "   Customer name, city, employment, income, tenure, DTI, life event\n"
    "   + Top 3 SHAP drivers (what's driving their risk)\n"
    "   + Intervention type (payment_holiday, emi_restructuring, etc.)\n\n"
    "3. GROQ API CALL: HTTP POST to api.groq.com with:\n"
    "   model='llama-3.3-70b-versatile', temperature=0.7, max_tokens=300\n"
    "   Response returns in ~0.5 seconds.\n\n"
    "4. BRAND SAFETY: System prompts enforce:\n"
    "   NEVER say 'delinquency', 'default', 'risk score' to the customer\n"
    "   Always use ₹ symbol and Indian English\n"
    "   Focus on HELP, not collections\n\n"
    "5. FALLBACK: If Groq API is unavailable (no key, rate limit), we use "
    "pre-written templates with {name} placeholders — functional but not personalized."
)

add_tech(
    "23.12 Debezium (Change Data Capture)",
    
    "Debezium enables real-time data synchronization:\n"
    "• Captures INSERT/UPDATE/DELETE events from PostgreSQL without polling\n"
    "• Zero impact on database performance (reads WAL, not tables)\n"
    "• Provides full 'before' and 'after' images of changed rows\n"
    "• Critical for event-driven architecture — downstream systems react to changes instantly",
    
    "Alternatives: Database triggers (blocks writes, adds latency), "
    "Polling (wastes CPU, misses rapid changes), "
    "Application-level events (requires modifying every write path). "
    "Debezium is the standard for CDC in event-driven architectures.",
    
    "Debezium as a security camera watching a database:\n"
    "• Watches the PostgreSQL transaction log (WAL) 24/7\n"
    "• When a row is inserted/updated/deleted, it captures the change\n"
    "• Publishes the change event to a Kafka topic\n"
    "• Downstream consumers (analytics, Flink) react in real-time\n\n"
    "Input: PostgreSQL WAL stream → Output: Kafka events for every row change",
    
    "Debezium internals:\n\n"
    "1. CONNECTOR: We register a PostgreSQL connector via REST API:\n"
    "   POST /connectors with config: slot.name='pdi_slot', plugin.name='pgoutput', "
    "tables: 'public.customers, public.transactions, public.risk_scores'\n\n"
    "2. WAL READING: PostgreSQL writes every transaction to the Write-Ahead Log. "
    "Debezium connects as a replication client and reads the WAL using PostgreSQL's "
    "logical decoding protocol (pgoutput plugin).\n\n"
    "3. EVENT STRUCTURE: Each change event contains:\n"
    "   • op: 'c' (create), 'u' (update), 'd' (delete)\n"
    "   • before: row values before the change (for updates/deletes)\n"
    "   • after: row values after the change (for creates/updates)\n"
    "   • source: database name, table, LSN (Log Sequence Number), timestamp\n\n"
    "4. KAFKA TOPICS: Events are published to topics named with the prefix:\n"
    "   pdi.public.customers, pdi.public.transactions, pdi.public.risk_scores\n\n"
    "5. EXACTLY-ONCE: Debezium tracks its position in the WAL via the replication slot. "
    "If it crashes, it resumes from the last acknowledged LSN — no events are lost or duplicated."
)

add_tech(
    "23.13 n8n (Workflow Automation)",
    
    "n8n was chosen for the demo workflow because:\n"
    "• Visual workflow editor — judges can SEE the data flow in real-time\n"
    "• HTTP Request node — directly calls our scoring API\n"
    "• Code node — runs JavaScript for data transformation\n"
    "• IF node — branches based on risk score thresholds\n"
    "• Self-hosted — runs in Docker, no external dependencies\n"
    "• Real-time execution visualization — nodes flash green/red as they execute",
    
    "Alternatives: Airflow (code-only, no visual execution), "
    "Zapier (cloud-only, not self-hosted), "
    "Prefect (less visual). n8n is ideal for hackathon demos because judges can watch "
    "the workflow execute in real-time.",
    
    "n8n as a flowchart that runs itself:\n"
    "• Every 10 seconds: generate a transaction → score the customer → check if high risk → send alert\n"
    "• Judges see each node light up as it executes\n"
    "• High-risk = red path (alert), Low-risk = green path (log)\n\n"
    "Input: Timer trigger → Output: Visual execution + notification",
    
    "n8n in our demo:\n\n"
    "1. WORKFLOW (n8n/pdi_demo_workflow.json): 6 nodes:\n"
    "   Schedule (every 10s) → Code (generate transaction) → HTTP Request (POST /score) "
    "   → IF (risk_score >= 0.5) → Code (format alert) or Code (log stable)\n\n"
    "2. EXECUTION ENGINE: n8n processes nodes sequentially. Each node's output becomes "
    "the next node's input. The IF node creates two branches based on the risk score.\n\n"
    "3. IMPORT: Judges import the JSON file via the n8n UI (⋮ → Import from file). "
    "Toggle the workflow active, and watch it execute every 10 seconds.\n\n"
    "4. DOCKER CONFIG: Runs as pdi-n8n container, port 5678, with N8N_BASIC_AUTH enabled "
    "for login security."
)

add_tech(
    "23.14 Grafana + Prometheus (Monitoring)",
    
    "Combined for production-grade observability:\n"
    "• Prometheus: pull-based metrics collection — FastAPI exposes /metrics endpoint\n"
    "• Grafana: real-time dashboards with alerting — visualizes latency, throughput, errors\n"
    "• prometheus-fastapi-instrumentator: auto-instruments every REST endpoint\n"
    "• Critical for production — detect API degradation before customers are affected",
    
    "Alternatives: Datadog (expensive SaaS), ELK Stack (log-focused not metrics-focused), "
    "CloudWatch (AWS-only). Grafana + Prometheus is the industry standard for open-source monitoring.",
    
    "Prometheus + Grafana as a car dashboard:\n"
    "• Prometheus = the sensors (speedometer, thermometer, fuel gauge)\n"
    "• Grafana = the dashboard display showing all gauges with alerts\n"
    "• If scoring latency exceeds 1 second, Grafana shows a red alert\n\n"
    "Input: Metrics from FastAPI → Output: Real-time charts + alerts",
    
    "Monitoring stack in our engine:\n\n"
    "1. PROMETHEUS METRICS: prometheus-fastapi-instrumentator adds:\n"
    "   • http_request_duration_seconds (histogram with p50, p95, p99)\n"
    "   • http_requests_total (counter by status code)\n"
    "   • Custom: pdi_scores_computed_total, pdi_interventions_sent_total\n\n"
    "2. SCRAPING: Prometheus (port 9090) scrapes FastAPI's /metrics every 15 seconds\n"
    "   and stores time-series data locally.\n\n"
    "3. GRAFANA DASHBOARDS: Pre-configured dashboards at port 3000 showing:\n"
    "   • Scoring latency over time (should stay < 100ms)\n"
    "   • Request rate (requests/second)\n"
    "   • Error rate (4xx, 5xx responses)\n"
    "   • Risk tier distribution over time"
)

add_tech(
    "23.15 Twilio (SMS & WhatsApp)",
    
    "Twilio for real notification delivery:\n"
    "• Industry-standard messaging API — used by Uber, Airbnb, Netflix\n"
    "• India-specific: supports +91 SMS delivery + DND handling\n"
    "• WhatsApp Business API — direct WhatsApp messages to customers\n"
    "• Delivery receipts — track if SMS was actually delivered\n"
    "• Free trial — sufficient for hackathon demo",
    
    "Alternatives: AWS SNS (complex setup), MSG91 (India-specific, less WhatsApp support), "
    "direct carrier APIs (requires separate integration per carrier). "
    "Twilio provides unified SMS + WhatsApp in one API.",
    
    "Twilio as a postal service:\n"
    "• We give it: phone number + message text\n"
    "• It delivers the SMS/WhatsApp message\n"
    "• It tells us: delivered, failed, or pending\n\n"
    "Input: Phone number + message → Output: Delivery status + SID",
    
    "Twilio in our notification dispatcher:\n\n"
    "1. SMS: client.messages.create(body=message, from_=TWILIO_FROM, to=phone)\n"
    "   Returns message SID (unique delivery tracking ID)\n"
    "   We store this SID in the notifications table for audit\n\n"
    "2. WHATSAPP: Same API but with 'whatsapp:' prefix on phone numbers\n"
    "   from_='whatsapp:+14155238886' (Twilio sandbox number)\n"
    "   to='whatsapp:+91XXXXXXXXXX'\n\n"
    "3. SIMULATION MODE: If TWILIO_ACCOUNT_SID is empty in .env, we log "
    "the notification to the DB without calling Twilio. This allows full "
    "demo without actual SMS costs.\n\n"
    "4. COOLDOWN: Before calling Twilio, we check Redis: was this customer "
    "contacted on this channel in the last 24 hours? If yes, skip."
)

add_tech(
    "23.16 MLflow (Experiment Tracking)",
    
    "MLflow tracks model experiments and versions:\n"
    "• Logs metrics (AUC, F1, accuracy) for every training run\n"
    "• Tracks hyperparameters — can compare 100 experiment configs side by side\n"
    "• Model versioning — knows which model is 'champion' (production) vs 'challenger'\n"
    "• Artifact storage — stores trained model files with metadata",
    
    "Alternatives: Weights & Biases (SaaS, costs money), "
    "TensorBoard (TensorFlow-only), Neptune (SaaS). "
    "MLflow is the only open-source, self-hosted experiment tracker with a model registry.",
    
    "MLflow as a lab notebook:\n"
    "• Every experiment (training run) is logged with its results\n"
    "• You can compare: 'Run #5 with lr=0.05 got 0.896 vs Run #3 with lr=0.1 got 0.872'\n"
    "• The best model is marked as 'champion' and used for production scoring\n\n"
    "Input: Metrics + params → Output: Comparison dashboard at localhost:5000",
    
    "MLflow in our pipeline:\n\n"
    "1. LOGGING (ml/mlflow_registry.py):\n"
    "   mlflow.log_params({'n_estimators': 300, 'max_depth': 6, ...})\n"
    "   mlflow.log_metrics({'auc': 0.896, 'f1': 0.82, ...})\n"
    "   mlflow.xgboost.log_model(model, 'xgboost_model')\n\n"
    "2. EXPERIMENT: All runs are grouped under experiment 'pdi_delinquency_prediction'\n\n"
    "3. BACKEND: MLflow stores metadata in PostgreSQL (mlflow_db database) and "
    "artifacts on local filesystem. Accessible via Web UI at port 5000.\n\n"
    "4. MODEL REGISTRY: After training, the ensemble is registered and the best version "
    "is promoted to 'Production' stage. Scoring service loads the production model."
)

# SAVE
doc.save(doc_path)
print(f"✅ DOCX updated with Technology Deep Dive: {doc_path}")
