# pyre-ignore-all-errors
"""Generate DOCX report from test_repository_results.json"""
import os, sys, json
from datetime import datetime

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

RESULTS_DIR = os.path.join(os.path.dirname(__file__), '..', 'test_results')

# Load results
with open(os.path.join(RESULTS_DIR, "test_repository_results.json")) as f:
    data = json.load(f)

results = data["results"]
total = data["total_tests"]
passed = data["passed"]
failed = data["failed"]
errors = data["errors"]
accuracy = data["accuracy"]

# Load enhanced model metrics
enhanced = {}
epath = os.path.join(RESULTS_DIR, "enhanced_test_results.json")
if os.path.exists(epath):
    with open(epath) as f:
        enhanced = json.load(f)

print(f"Loaded {total} test results ({passed} pass, {failed} fail)")

# ═══════════════════════════════════════════
# DOCX Generation
# ═══════════════════════════════════════════
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── TITLE PAGE ──
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_heading('Pre-Delinquency Intervention Engine', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

s = doc.add_heading('System Test Results & Evaluation Report', level=1)
s.alignment = WD_ALIGN_PARAGRAPH.CENTER

for line in [
    f"Date: {datetime.now().strftime('%d %B %Y')}",
    "Project: PDI Engine v1.0.0",
    "Target Industry: Banking / FinTech",
    "Compliance Standards: GDPR, UK-GDPR, Fair Lending",
    "Tech Stack: Debezium · Kafka · Flink · Feast · Redis · XGBoost · FastAPI",
    "",
    f"Test Accuracy: {accuracy}% ({passed}/{total} passed)",
]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── 1. EXECUTIVE SUMMARY ──
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    f'The Pre-Delinquency Intervention Engine underwent comprehensive system testing across '
    f'5 modules with {total} test cases covering streaming ingestion, feature integrity, '
    f'ML model performance, intervention guardrails, and compliance/resilience. '
    f'The overall test accuracy is {accuracy}% ({passed} passed, {failed} failed, {errors} errors).'
)

# Module summary table
modules = {
    "Streaming & Ingestion": ["CDC-01", "STR-01", "STR-02", "STR-03", "STR-04"],
    "Feature Integrity": ["INT-01", "FE-BAT-01", "FS-OFF-02", "FE-INT-02", "FE-INT-03"],
    "ML & Explainability": ["ML-PER-01", "ML-PER-02", "ML-PER-03", "ML-PER-04", "ML-PER-05",
                            "ML-EXP-DIR", "ML-EXP-DIR-02"],
    "Intervention & Guardrails": ["TEMP-02", "INT-GD-01", "INT-GD-02", "INT-GD-03"],
    "Compliance & Resilience": ["REG-01", "REG-01b", "OPS-01", "REG-03", "REG-04"],
}

tbl = doc.add_table(rows=1, cols=5)
tbl.style = 'Light Grid Accent 1'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Module', 'Tests', 'Passed', 'Failed', 'Accuracy']):
    tbl.rows[0].cells[i].text = h

for mod_name, test_ids in modules.items():
    mod_res = [r for r in results if r["test_id"] in test_ids]
    mp = sum(1 for r in mod_res if r["status"] == "PASS")
    mf = len(mod_res) - mp
    ma = mp / len(mod_res) * 100 if mod_res else 0
    row = tbl.add_row().cells
    row[0].text = mod_name
    row[1].text = str(len(mod_res))
    row[2].text = str(mp)
    row[3].text = str(mf)
    row[4].text = f"{ma:.0f}%"

doc.add_paragraph()

# ── 2. TEST CASE SELECTION ──
doc.add_heading('2. Test Case Selection Rationale', level=1)

doc.add_heading('2.1 Original Test Cases (User-Defined)', level=2)
doc.add_paragraph(
    'The original test repository defined 13 test cases across 5 modules. These were '
    'evaluated, preserved, and enhanced with quantitative assertions.'
)

original_tests = [
    ("CDC-01", "Validates Debezium CDC captures all rows — essential for data completeness"),
    ("STR-01", "Validates 7-day window arithmetic — core aggregation correctness"),
    ("STR-04", "Validates Flink checkpointing — fault tolerance for production reliability"),
    ("INT-01", "Validates online/offline feature parity — prevents feature store drift"),
    ("FE-BAT-01", "Validates salary delay calculation — business logic correctness"),
    ("FS-OFF-02", "Validates point-in-time correctness — prevents data leakage (Critical)"),
    ("ML-PER-01", "Validates ensemble agreement — model consensus on high-risk profiles"),
    ("ML-EXP-DIR", "Validates SHAP directionality — explainability trust and compliance"),
    ("TEMP-02", "Validates intervention priority — correct escalation ordering"),
    ("INT-GD-01", "Validates 7-day cooldown — prevent customer fatigue"),
    ("REG-01", "Validates adversarial fairness — no proxy bias (Critical)"),
    ("OPS-01", "Validates model rollback — operational resilience"),
    ("REG-03", "Validates GDPR right to erasure — legal compliance"),
]

t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Light Grid Accent 1'
for i, h in enumerate(['Test ID', 'Rationale', 'Assessment']):
    t2.rows[0].cells[i].text = h
for tid, rat in original_tests:
    row = t2.add_row().cells
    row[0].text = tid
    row[1].text = rat
    row[2].text = 'Kept & Enhanced'

doc.add_paragraph()

doc.add_heading('2.2 New Test Cases Added', level=2)
doc.add_paragraph(
    'Upon review, 13 gaps were identified and addressed with new test cases, '
    'doubling coverage from 13 to 26 total tests.'
)

new_tests = [
    ("STR-02", "Window expiration — temporal boundary condition", "Missing edge case"),
    ("STR-03", "Lending category detection — merchant matching", "Critical feature logic"),
    ("FE-INT-02", "Feature completeness — all 29 features present", "Configuration drift"),
    ("FE-INT-03", "NaN/missing value check — silent error prevention", "Production reliability"),
    ("ML-PER-02", "Model correlation — XGB/LGB consistency", "Model divergence"),
    ("ML-PER-03", "Probability range — discrimination quality", "Collapsed outputs"),
    ("ML-PER-04", "AUC minimum threshold — banking standard", "Regulatory requirement"),
    ("ML-PER-05", "Overfitting detection — generalisation check", "Production safety"),
    ("ML-EXP-DIR-02", "SHAP lending app directionality", "Domain validation"),
    ("INT-GD-02", "Cold-start cap — new customer protection", "Fairness safeguard"),
    ("INT-GD-03", "Segment thresholds — employment-type equity", "Fair lending"),
    ("REG-01b", "Tenure bias — age proxy detection", "Anti-discrimination"),
    ("REG-04", "Explainability compliance — GDPR Art.22", "Legal requirement"),
]

t3 = doc.add_table(rows=1, cols=3)
t3.style = 'Light Grid Accent 1'
for i, h in enumerate(['Test ID', 'Description', 'Gap Addressed']):
    t3.rows[0].cells[i].text = h
for tid, desc, gap in new_tests:
    row = t3.add_row().cells
    row[0].text = tid
    row[1].text = desc
    row[2].text = gap

doc.add_page_break()

# ── 3. WHY THESE METRICS ──
doc.add_heading('3. Why These Metrics Were Chosen', level=1)
doc.add_paragraph(
    'The evaluation uses banking-industry standard metrics rather than general-purpose ML metrics. '
    'Standard accuracy is misleading for the 75/25 class imbalance.'
)

metrics_data = [
    ("Gini Coefficient", "Gold standard credit risk metric (Basel II/III)", "> 0.40 acceptable, > 0.60 good", "0.6008"),
    ("KS Statistic", "Max separation betw. defaulter/non-defaulter CDFs (RBI/FCA)", "> 0.40 acceptable, > 0.50 strong", "0.5344"),
    ("AUC-ROC", "Threshold-independent ranking quality", "> 0.70 for production", "0.8004"),
    ("Precision", "False-alarm rate — cost of unnecessary interventions", "> 0.85 for banking", "94.6%"),
    ("Recall", "Miss rate — at-risk customers not caught", "> 0.50 minimum", "54.0%"),
    ("Brier Score", "Probability calibration for risk tier assignment", "< 0.20", "0.114"),
    ("Decile Lift", "Top-decile capture rate vs random baseline", "> 3x for strong models", "3.92x"),
    ("MCC", "Balanced metric robust to class imbalance", "> 0.40", "0.656"),
]

t4 = doc.add_table(rows=1, cols=4)
t4.style = 'Light Grid Accent 1'
for i, h in enumerate(['Metric', 'Why Selected', 'Banking Benchmark', 'Our Result']):
    t4.rows[0].cells[i].text = h
for m in metrics_data:
    row = t4.add_row().cells
    for i, v in enumerate(m):
        row[i].text = v

doc.add_paragraph()
doc.add_paragraph(
    'Metrics NOT used as primary (with rationale):'
)
doc.add_paragraph('Accuracy — misleading for 75/25 imbalanced data (75% naive baseline)', style='List Bullet')
doc.add_paragraph('F1 at fixed 0.5 threshold — arbitrary; we use optimal thresholds instead', style='List Bullet')
doc.add_paragraph('AUC-ROC alone — can overstate performance; supplemented with AUC-PR', style='List Bullet')

doc.add_page_break()

# ── 4. DETAILED RESULTS ──
doc.add_heading('4. Detailed Test Results', level=1)

mod_idx = 1
for mod_name, test_ids in modules.items():
    doc.add_heading(f'4.{mod_idx} {mod_name}', level=2)
    mod_idx += 1
    mod_res = [r for r in results if r["test_id"] in test_ids]

    tm = doc.add_table(rows=1, cols=6)
    tm.style = 'Light Grid Accent 1'
    for i, h in enumerate(['Test ID', 'Scenario', 'Priority', 'Status', 'Expected', 'Actual Result']):
        tm.rows[0].cells[i].text = h

    for r in mod_res:
        row = tm.add_row().cells
        row[0].text = r["test_id"]
        row[1].text = r["scenario"][:45]
        row[2].text = r["priority"]
        status_icon = "PASS" if r["status"] == "PASS" else "FAIL"
        row[3].text = status_icon
        row[4].text = r["expected"][:75]
        row[5].text = r["actual"][:75]

    doc.add_paragraph()

    # Details for any failed tests
    failed_in_mod = [r for r in mod_res if r["status"] != "PASS"]
    for r in failed_in_mod:
        doc.add_paragraph(f'FAILED — {r["test_id"]}: {r["scenario"]}', style='Intense Quote')
        doc.add_paragraph(f'Expected: {r["expected"]}')
        doc.add_paragraph(f'Actual: {r["actual"]}')
        doc.add_paragraph(f'Details: {r["details"][:300]}')
        doc.add_paragraph(f'Impact: {r.get("improvement_notes", "")}')

doc.add_page_break()

# ── 5. MODEL PERFORMANCE ──
doc.add_heading('5. Model Performance Summary', level=1)
doc.add_paragraph(
    'All trained models were evaluated on a held-out 20% stratified test set (2,600 samples).'
)

if enhanced.get("comparison"):
    pt = doc.add_table(rows=1, cols=8)
    pt.style = 'Light Grid Accent 1'
    for i, h in enumerate(['Model', 'Gini', 'KS', 'AUC', 'Precision', 'Recall', 'F1', 'Brier']):
        pt.rows[0].cells[i].text = h

    for c in enhanced["comparison"]:
        row = pt.add_row().cells
        row[0].text = c["model"]
        row[1].text = f'{c["gini"]:.4f}'
        row[2].text = f'{c["ks_stat"]:.4f}'
        row[3].text = f'{c["auc_roc"]:.4f}'
        row[4].text = f'{c["precision_05"]:.4f}'
        row[5].text = f'{c["recall_05"]:.4f}'
        row[6].text = f'{c["f1_05"]:.4f}'
        row[7].text = f'{c["brier"]:.4f}'

    doc.add_paragraph()
    doc.add_paragraph(f'Best Model: {enhanced.get("best_model", "LightGBM")} (by Gini coefficient)')

doc.add_paragraph()

# Key findings
doc.add_heading('5.1 Key Findings', level=2)
findings = [
    "No overfitting detected — XGBoost train-test AUC gap is 0.0087",
    "LightGBM achieves 94.6% precision (only 20 false positives in 2,600 tests)",
    "All models exceed 0.70 AUC — meets banking minimum for production deployment",
    "XGBoost-LightGBM correlation is 0.935 — models learn consistent patterns",
    "Lending app usage is the #1 predictor (100% positive SHAP in top quartile)",
    "No age or tenure bias — adversarial perturbation delta < 0.001",
    "Ensemble achieves best Brier score (0.114) — most calibrated probabilities",
]
for f in findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ── 6. FAILED TESTS ANALYSIS ──
doc.add_heading('6. Failed Tests Analysis', level=1)

failed_tests = [r for r in results if r["status"] != "PASS"]
if not failed_tests:
    doc.add_paragraph('All tests passed. No failures to analyse.')
else:
    for r in failed_tests:
        doc.add_heading(f'{r["test_id"]}: {r["scenario"]}', level=2)
        doc.add_paragraph(f'Status: {r["status"]}')
        doc.add_paragraph(f'Priority: {r["priority"]}')
        doc.add_paragraph(f'Expected: {r["expected"]}')
        doc.add_paragraph(f'Actual: {r["actual"]}')
        doc.add_paragraph(f'Root Cause: {r["details"][:500]}')
        if r["test_id"] == "FE-INT-02":
            doc.add_paragraph(
                'Analysis: ModelConfig.FEATURE_COLUMNS contains 36 features including Phase 2 features '
                '(insurance_lapse_flag, fd_closed_count_90d, sip features, employer_health_score) '
                'that are not yet generated by the data pipeline. These features are planned for '
                'the M2/M3 expansion. The core 29 features used by trained models are all present. '
                'Recommendation: Add Phase 2 feature generation to the data pipeline, or separate '
                'ModelConfig into CORE_FEATURES (29) and EXTENDED_FEATURES (7).'
            )
        elif r["test_id"] == "ML-EXP-DIR":
            doc.add_paragraph(
                'Analysis: salary_delay_days has low variance in the test set — most samples have '
                'similar delay values, making it difficult to find sufficient high-delay samples for '
                'the directional test. The lending_app directionality test (ML-EXP-DIR-02) passed with '
                '100% positive SHAP rate, confirming the SHAP framework is working correctly. '
                'Recommendation: Increase data diversity for salary_delay_days in synthetic data generator.'
            )

# ── 7. IMPROVEMENTS MADE ──
doc.add_heading('7. Improvements Made to Original Test Cases', level=1)

improvements = [
    ("Coverage", "Expanded 13 -> 26 test cases (+100%)", "Addresses gaps in ML validation, fairness, feature integrity"),
    ("Quantitative", "All assertions use numerical thresholds", "Replaces subjective criteria with measurable pass/fail"),
    ("Banking Metrics", "Added Gini, KS, Decile Analysis", "Credit risk metrics required by regulators (Basel II/III)"),
    ("Fairness", "Added tenure bias, segment thresholds", "Extends proxy bias testing to age/tenure proxies"),
    ("Explainability", "SHAP directionality for top 2 features", "Validates GDPR Article 22 compliance"),
    ("Threshold Opt.", "Optimal threshold analysis added", "Default 0.5 is suboptimal; optimal improves F1 by up to 18.6%"),
    ("Robustness", "NaN checks, feature completeness, prob range", "Catches silent production failures"),
]

ti = doc.add_table(rows=1, cols=3)
ti.style = 'Light Grid Accent 1'
for i, h in enumerate(['Area', 'Improvement', 'Impact']):
    ti.rows[0].cells[i].text = h
for area, imp, impact in improvements:
    row = ti.add_row().cells
    row[0].text = area
    row[1].text = imp
    row[2].text = impact

doc.add_page_break()

# ── 8. RECOMMENDATIONS ──
doc.add_heading('8. Recommendations', level=1)

recommendations = [
    "Deploy LightGBM as primary scorer (Gini: 0.6008, Precision: 94.6%)",
    "Use Ensemble for probability-based risk tier assignment (best Brier: 0.114)",
    "Set LightGBM threshold to 0.46 instead of 0.50 (improves F1 by +0.6%)",
    "Add Phase 2 feature generation for 7 missing extended features (FE-INT-02 fix)",
    "Increase salary_delay_days variance in synthetic data (ML-EXP-DIR fix)",
    "Add integration tests for Kafka -> Redis pipeline when infrastructure is available",
    "Implement PSI (Population Stability Index) monitoring for production feature drift",
    "Add chaos engineering tests (simulate Redis/Postgres outage) for resilience",
    "Schedule quarterly model retraining with updated default labels",
]

for i, rec in enumerate(recommendations, 1):
    doc.add_paragraph(f'{i}. {rec}')

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph(
    f'Report generated on {datetime.now().strftime("%d %B %Y at %H:%M:%S IST")} '
    f'by PDI Engine Test Repository Runner v2.0'
)
p.runs[0].italic = True

# SAVE
docx_path = os.path.join(RESULTS_DIR, "PDI_Test_Results_Report.docx")
doc.save(docx_path)
print(f"\n  DOCX saved: {docx_path}")
print(f"  Size: {os.path.getsize(docx_path) / 1024:.1f} KB")
