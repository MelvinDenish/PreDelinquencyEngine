# pyre-ignore-all-errors
"""
Generate SYSTEM_ARCHITECTURE.docx — Comprehensive system documentation
for the Pre-Delinquency Intervention Engine hackathon submission.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ─── Styles ───
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title page
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("PRE-DELINQUENCY INTERVENTION ENGINE")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 48, 135)  # Barclays blue

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Complete System Architecture & Implementation Guide")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Team Solace — College of Engineering, Guindy\n").font.size = Pt(12)
meta.add_run("Barclays Hackathon 2026\n").font.size = Pt(12)
meta.add_run("\nConfidential").font.color.rgb = RGBColor(180, 0, 0)

doc.add_page_break()

# ════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Black-Box View: What the System Does",
    "3. System Architecture Overview",
    "4. Data Layer",
    "   4.1. Customer Profile Generator",
    "   4.2. Transaction Generator",
    "   4.3. PostgreSQL Schema",
    "5. Streaming Pipeline (Apache Flink + Kafka)",
    "6. Batch Processing (Apache Spark)",
    "7. Feature Store (Feast)",
    "8. ML Models",
    "   8.1. XGBoost (Gradient Boosting)",
    "   8.2. LightGBM",
    "   8.3. LSTM (Deep Learning)",
    "   8.4. Ensemble Scoring",
    "9. Explainability (SHAP + LIME)",
    "10. Fairness Audit (Fairlearn + AIF360)",
    "11. Scoring Service (FastAPI)",
    "12. Notification Dispatcher",
    "   12.1. Channel Routing Matrix",
    "   12.2. Email (SMTP)",
    "   12.3. SMS & WhatsApp (Twilio)",
    "   12.4. RM Call Assignment",
    "   12.5. Collector Escalation",
    "13. GenAI Message Generation (Groq LLM)",
    "14. Rules Engine & Intervention Logic",
    "15. Dashboard (Plotly Dash)",
    "16. Workflow Orchestration (n8n + Airflow)",
    "17. Monitoring (Grafana + Prometheus)",
    "18. CDC Pipeline (Debezium)",
    "19. Deployment Architecture (Docker + K8s)",
    "20. Technology Stack Summary",
    "21. Data Flow Diagram",
    "22. API Reference",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "The Pre-Delinquency Intervention Engine (PDI Engine) is a production-grade, "
    "end-to-end machine learning platform designed to predict which bank customers "
    "are likely to become delinquent on their loan payments BEFORE it happens. "
    "Rather than waiting for missed payments (reactive collections), the system "
    "detects early warning signals in real-time transaction data and triggers "
    "proactive, personalized interventions through multiple channels."
)
doc.add_paragraph(
    "The engine processes 37,000+ transactions across 1,000 customers, computes "
    "40+ behavioral features using both real-time (Flink) and batch (Spark) pipelines, "
    "and uses a 3-model ensemble (XGBoost + LightGBM + LSTM) achieving an AUC of 0.896 "
    "to score customers into risk tiers. Based on the risk tier, the system automatically "
    "routes personalized GenAI-crafted messages through SMS, Email, WhatsApp, RM calls, "
    "or collector assignments — all tracked in a production notification database."
)

doc.add_heading("Key Outcomes", level=2)
outcomes = [
    "Ensemble AUC: 0.896 — accurately identifies at-risk customers",
    "5 notification channels: SMS, Email, WhatsApp, RM Call, Collector Assignment",
    "Real-time scoring: < 100ms per customer via FastAPI",
    "16 Docker containers: fully containerized microservices architecture",
    "Explainable AI: SHAP + LIME for every prediction",
    "Fairness audited: Fairlearn + AIF360 bias checks across gender, age, region",
    "GenAI-powered: Groq LLM generates personalized intervention messages",
]
for o in outcomes:
    doc.add_paragraph(o, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════
# 2. BLACK-BOX VIEW
# ════════════════════════════════════════
doc.add_heading("2. Black-Box View: What the System Does", level=1)
doc.add_paragraph(
    "From a user's perspective (without knowing the internals), the system works as follows:"
)

doc.add_heading("Inputs", level=2)
inputs = [
    ("Customer Data", "Demographics, income, employment, credit score, loan details"),
    ("Transaction Stream", "Real-time UPI, ATM, bill payments, EMI auto-debits, salary credits"),
    ("Account Balances", "Daily savings and current account balance snapshots"),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Input"
hdr[1].text = "Description"
for name, desc in inputs:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = desc

doc.add_heading("Outputs", level=2)
outputs = [
    ("Risk Score", "0.0–1.0 probability of delinquency for each customer"),
    ("Risk Tier", "stable (< 0.5), watch (0.5–0.7), critical (≥ 0.7)"),
    ("SHAP Drivers", "Top 5 features driving the risk score (explainable)"),
    ("SMS Notification", "160-char empathetic nudge sent via Twilio"),
    ("Email", "Branded HTML email with financial wellness offers"),
    ("WhatsApp Message", "Friendly message with CTA"),
    ("RM Call Task", "Callback task with GenAI call script for relationship manager"),
    ("Collector Assignment", "Case brief with restructuring offer for collections team"),
    ("Dashboard", "Real-time risk monitoring with charts and distributions"),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Output"
hdr[1].text = "Description"
for name, desc in outputs:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = desc

doc.add_heading("Black-Box Flow", level=2)
doc.add_paragraph(
    "Customer transacts → System detects stress signals → ML scores risk → "
    "If risk ≥ threshold → GenAI generates personalized message → "
    "System sends via appropriate channel → RM/Collector assigned if severe → "
    "All tracked in notification database"
)

doc.add_page_break()

# ════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ════════════════════════════════════════
doc.add_heading("3. System Architecture Overview", level=1)
doc.add_paragraph(
    "The system follows a Lambda Architecture pattern combining real-time (speed layer) "
    "and batch (batch layer) processing with a serving layer for low-latency scoring."
)

doc.add_heading("Architecture Layers", level=2)
layers = [
    ("Ingestion Layer", "Debezium CDC captures PostgreSQL changes → Kafka topics. "
     "Transaction generator publishes directly to Kafka."),
    ("Speed Layer (Real-time)", "Apache Flink (or local Python consumer) processes "
     "Kafka streams, computes 7-day and 30-day rolling features like discretionary "
     "spending, ATM withdrawal counts, lending app usage, failed auto-debits."),
    ("Batch Layer", "Apache Spark (or local Pandas fallback) computes historical "
     "features: salary delay patterns, utility payment delays, spend volatility, "
     "discretionary spend trends over 3-month windows."),
    ("Feature Store", "Feast materializes features from both layers into a unified "
     "feature view for consistent training and serving."),
    ("ML Layer", "Three models trained: XGBoost (tabular, gradient boosting), "
     "LightGBM (tabular, light gradient boosting), LSTM (temporal sequences). "
     "Combined via weighted ensemble (0.6 × XGBoost + 0.4 × LSTM)."),
    ("Serving Layer", "FastAPI scoring service loads all three models at startup. "
     "Receives customer_id, retrieves features from PostgreSQL, scores via ensemble, "
     "computes SHAP explanations, stores results, and returns response in < 100ms."),
    ("Intervention Layer", "Rules engine determines intervention type from SHAP drivers. "
     "Notification dispatcher routes to SMS/Email/WhatsApp/RM/Collector based on risk tier. "
     "GenAI (Groq LLM) generates personalized messages for each channel."),
    ("Monitoring Layer", "Prometheus collects metrics from scoring service. "
     "Grafana dashboards visualize latency, throughput, error rates. "
     "Drift detector monitors feature and prediction distributions."),
]
for name, desc in layers:
    doc.add_heading(name, level=3)
    doc.add_paragraph(desc)

doc.add_page_break()

# ════════════════════════════════════════
# 4. DATA LAYER
# ════════════════════════════════════════
doc.add_heading("4. Data Layer", level=1)

doc.add_heading("4.1. Customer Profile Generator", level=2)
doc.add_paragraph(
    "File: data_generator/customer_generator.py\n\n"
    "Generates production-grade customer profiles with:"
)
cust_features = [
    "9 employment types: salaried_private, salaried_govt, self_employed, professional, "
    "business_owner, contract_worker, gig_worker, freelancer, retired",
    "Industry sectors mapped to employment (IT/ITES, BFSI, Manufacturing, Healthcare, etc.)",
    "7 income brackets: EWS to Ultra HNI",
    "12 life events as stress triggers: job_loss, medical_emergency, salary_cut, "
    "divorce, business_failure, wedding_expense, etc.",
    "DTI (Debt-to-Income) ratio calculated from actual product holdings",
    "CIBIL-realistic credit scores (300–900) adjusted by employment type and tenure",
    "Contact details: phone (+91XXXXXXXXXX), email with realistic domains",
    "Risk segments: new_to_bank, stable_core, high_value, vulnerable",
    "Channel preferences weighted by age and income (older → SMS/RM, younger → app/WhatsApp)",
]
for f in cust_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading("4.2. Transaction Generator", level=2)
doc.add_paragraph(
    "File: data_generator/transaction_generator.py\n\n"
    "Generates 6 months of transaction history per customer with:"
)
txn_features = [
    "18 merchant categories: grocery, dining, entertainment, utility, rent, "
    "lending_app, gambling, crypto_exchange, payday_lender, etc.",
    "Salary credits on expected day (with delays for stressed customers)",
    "EMI auto-debit payments with failure simulation (40% fail rate for stressed)",
    "ATM withdrawals with stress-aware patterns (6–10/month for stressed vs 2–3 normal)",
    "Stress progression: intensity ramps up in last 2–3 months",
    "Income-scaled amounts: spending proportional to salary",
    "Channel-realistic: UPI for grocery/dining, netbanking for bills, auto for EMI",
]
for f in txn_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading("4.3. PostgreSQL Schema", level=2)
doc.add_paragraph(
    "File: init_db.sql\n\n"
    "12 tables in pdi_db database:"
)
tables_data = [
    ("customers", "Customer profiles with demographics, salary, products, channel preferences"),
    ("transactions", "All transaction records with merchant category, amount, status"),
    ("account_balances", "Daily balance snapshots (savings + current)"),
    ("streaming_features", "Real-time features from Flink (7d + 30d windows)"),
    ("batch_features", "Historical features from Spark (salary delay, spend trends)"),
    ("risk_scores", "Scored results with ensemble, XGBoost, LSTM scores + SHAP"),
    ("interventions", "Intervention records with type, channel, trigger reason"),
    ("feedback_events", "Outcome tracking (paid, restructured, defaulted)"),
    ("notifications", "Audit log of all notification attempts across all channels"),
    ("rm_tasks", "RM callback tasks with call scripts, priority, due dates"),
    ("collector_assignments", "Collection cases with restructuring offers"),
    ("model_registry", "Model versions, metrics, champion flag"),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Table"
hdr[1].text = "Purpose"
for name, desc in tables_data:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = desc

doc.add_page_break()

# ════════════════════════════════════════
# 5-7. PROCESSING PIPELINES
# ════════════════════════════════════════
doc.add_heading("5. Streaming Pipeline (Apache Flink + Kafka)", level=1)
doc.add_paragraph(
    "File: stream_processing/flink_job.py\n\n"
    "The streaming layer consumes transactions from Kafka in real-time and computes "
    "rolling window features:"
)
streaming_feats = [
    "discretionary_spend_7d / 30d: Total spending on dining, entertainment, clothing, luxury, travel",
    "atm_withdrawals_count_7d / 30d: Number of ATM withdrawals (stress signal)",
    "lending_app_txn_count_7d / 30d: Transactions to lending apps (high-risk signal)",
    "weighted_lending_risk_7d: Lending count × 0.85 risk weight",
    "failed_autodebits_count_7d / 30d: Failed EMI/auto-debit attempts",
    "total_spend_7d / 30d: Total debit transaction amount",
    "txn_count_7d / 30d: Total transaction count",
    "avg_txn_amount_7d: Average transaction size",
    "max_txn_amount_7d: Largest single transaction (unusual spike detection)",
    "savings_balance_pct_change_7d: Savings drawdown rate",
]
for f in streaming_feats:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph(
    "\nHow it works internally:\n"
    "1. Kafka consumer reads from 'pdi-transactions' topic\n"
    "2. Each transaction is parsed and grouped by customer_id\n"
    "3. Rolling windows (7-day and 30-day) are maintained in memory\n"
    "4. Features are computed per customer and upserted to streaming_features table\n"
    "5. Falls back to local Python consumer if Flink cluster is unavailable"
)

doc.add_heading("6. Batch Processing (Apache Spark)", level=1)
doc.add_paragraph(
    "File: batch_processing/spark_jobs.py (with local Pandas fallback in main.py)\n\n"
    "Computes historical features that require longer time windows:"
)
batch_feats = [
    "salary_delay_days: Average days late for salary credit vs expected day",
    "utility_payment_delay_avg: Average days late for utility bill payments",
    "discretionary_spend_trend: Ratio of current vs previous period discretionary spending",
    "avg_monthly_spend_3m: Average monthly debit spend over 3 months",
    "spend_volatility_3m: Standard deviation / mean of monthly spending (variability)",
    "product_count, has_credit_card, has_personal_loan, has_mortgage: Product flags",
    "Demographics: credit_score, age, tenure_months, income_bracket, region, gender",
]
for f in batch_feats:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading("7. Feature Store (Feast)", level=1)
doc.add_paragraph(
    "File: feature_store/materialize.py\n\n"
    "Feast provides a unified feature registry that ensures consistency between "
    "training and serving:\n"
    "1. Features exported from PostgreSQL to Parquet files\n"
    "2. feast apply registers feature views\n"
    "3. feast materialize loads features into online store\n"
    "4. Scoring service retrieves features via Feast get_online_features()"
)

doc.add_page_break()

# ════════════════════════════════════════
# 8. ML MODELS
# ════════════════════════════════════════
doc.add_heading("8. ML Models", level=1)
doc.add_paragraph(
    "File: ml/train.py orchestrates the 8-step training pipeline.\n\n"
    "Three complementary models are trained:"
)

doc.add_heading("8.1. XGBoost (Gradient Boosting)", level=2)
doc.add_paragraph(
    "File: ml/xgboost_model.py\n\n"
    "XGBoost is the primary model for tabular features. It excels at capturing "
    "non-linear relationships between features like salary delay, ATM withdrawal "
    "patterns, and lending app usage.\n\n"
    "Architecture:\n"
    "• Objective: binary:logistic (probability of delinquency)\n"
    "• Evaluation: AUC-ROC\n"
    "• Regularization: L1 + L2 to prevent overfitting\n"
    "• 5-fold stratified cross-validation\n"
    "• Early stopping on validation AUC\n\n"
    "Key hyperparameters:\n"
    "• n_estimators: 300 (with early stopping)\n"
    "• max_depth: 6\n"
    "• learning_rate: 0.05\n"
    "• subsample: 0.8\n"
    "• colsample_bytree: 0.8\n"
    "• scale_pos_weight: auto-calculated for class imbalance"
)

doc.add_heading("8.2. LightGBM", level=2)
doc.add_paragraph(
    "File: ml/lightgbm_model.py\n\n"
    "LightGBM provides a second opinion using histogram-based gradient boosting. "
    "It's faster than XGBoost and handles categorical features natively.\n\n"
    "Architecture:\n"
    "• Objective: binary (log loss)\n"
    "• Boosting: GOSS (Gradient-based One-Side Sampling)\n"
    "• 5-fold stratified cross-validation\n"
    "• Feature importance via split-based and gain-based methods\n\n"
    "Key hyperparameters:\n"
    "• n_estimators: 300\n"
    "• num_leaves: 31\n"
    "• learning_rate: 0.05\n"
    "• min_child_samples: 20\n"
    "• is_unbalance: True"
)

doc.add_heading("8.3. LSTM (Deep Learning)", level=2)
doc.add_paragraph(
    "File: ml/lstm_model.py\n\n"
    "The LSTM model captures temporal patterns in transaction sequences that "
    "tree-based models miss. It processes 30-day windows of daily feature vectors.\n\n"
    "Architecture:\n"
    "• Input: (batch_size, 30, num_features) — 30 days of daily features\n"
    "• 2 LSTM layers with 64 hidden units each\n"
    "• Dropout: 0.3 between layers\n"
    "• Fully connected output: hidden → 32 → 1 (sigmoid)\n"
    "• Optimizer: Adam (lr=0.001)\n"
    "• Loss: BCEWithLogitsLoss\n"
    "• Training: 30 epochs, batch_size=64\n\n"
    "Why LSTM matters:\n"
    "• Detects gradual deterioration patterns (spending creeping up over weeks)\n"
    "• Captures temporal dependencies (salary delay THIS month following ATM spike LAST week)\n"
    "• Complements tree models which treat features as static snapshots"
)

doc.add_heading("8.4. Ensemble Scoring", level=2)
doc.add_paragraph(
    "File: ml/ensemble.py\n\n"
    "The three models are combined via weighted average:\n\n"
    "    ensemble_score = 0.6 × XGBoost + 0.4 × LSTM\n"
    "    (LightGBM is used as XGBoost fallback and for feature importance validation)\n\n"
    "Risk Tier Classification:\n"
    "    • stable: score < 0.5 (no intervention needed)\n"
    "    • watch: 0.5 ≤ score < 0.7 (proactive nudge)\n"
    "    • critical: score ≥ 0.7 (urgent intervention required)\n\n"
    "Achieved Performance:\n"
    "    • XGBoost AUC: 0.896\n"
    "    • LightGBM AUC: 0.896\n"
    "    • LSTM AUC: 0.741\n"
    "    • Ensemble AUC: 0.896"
)

doc.add_page_break()

# ════════════════════════════════════════
# 9-10. EXPLAINABILITY & FAIRNESS
# ════════════════════════════════════════
doc.add_heading("9. Explainability (SHAP + LIME)", level=1)
doc.add_paragraph(
    "File: ml/explainability.py, ml/lime_explainer.py\n\n"
    "Every prediction includes explainability via two complementary methods:"
)

doc.add_heading("SHAP (SHapley Additive exPlanations)", level=3)
doc.add_paragraph(
    "• Uses TreeExplainer for XGBoost (exact Shapley values)\n"
    "• Returns top 5 features driving each prediction\n"
    "• Example: 'salary_delay_days pushed risk up by +0.18'\n"
    "• Used by rules engine to determine intervention TYPE\n"
    "  (salary delay → payment holiday offer, lending app usage → wellness check-in)"
)

doc.add_heading("LIME (Local Interpretable Model-agnostic Explanations)", level=3)
doc.add_paragraph(
    "• Perturbs input features and observes prediction changes\n"
    "• Provides human-readable explanations for non-technical stakeholders\n"
    "• Example: 'High risk because: ATM withdrawals increased 3x, "
    "lending app usage detected, savings dropped 35%'\n"
    "• Used in RM call scripts and collector briefs"
)

doc.add_heading("10. Fairness Audit (Fairlearn + AIF360)", level=1)
doc.add_paragraph(
    "File: ml/fairness.py\n\n"
    "The model is audited for bias across protected attributes:\n"
    "• Gender (Male vs Female)\n"
    "• Age groups (21-30, 31-45, 46-60, 60+)\n"
    "• Region (North, South, East, West, Central)\n\n"
    "Metrics checked:\n"
    "• Demographic Parity: equal positive prediction rates across groups\n"
    "• Equalized Odds: equal TPR and FPR across groups\n"
    "• Disparate Impact: ratio ≥ 0.8 (80% rule)\n\n"
    "Frameworks: Fairlearn (Microsoft) + AIF360 (IBM)"
)

doc.add_page_break()

# ════════════════════════════════════════
# 11. SCORING SERVICE
# ════════════════════════════════════════
doc.add_heading("11. Scoring Service (FastAPI)", level=1)
doc.add_paragraph(
    "File: scoring_service/app.py\n\n"
    "REST API that provides real-time scoring:"
)
endpoints = [
    ("POST /score", "Score a single customer by customer_id. Returns risk_score, "
     "risk_tier, SHAP top drivers, ensemble breakdown."),
    ("GET /health", "Health check — returns model load status"),
    ("GET /metrics", "Prometheus metrics endpoint"),
    ("GET /docs", "Swagger UI for interactive API testing"),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Endpoint"
hdr[1].text = "Description"
for ep, desc in endpoints:
    row = table.add_row().cells
    row[0].text = ep
    row[1].text = desc

doc.add_paragraph(
    "\nInternal flow for POST /score:\n"
    "1. Receive customer_id\n"
    "2. Fetch streaming_features + batch_features from PostgreSQL\n"
    "3. Merge into single feature vector (40+ features)\n"
    "4. Score via XGBoost → probability\n"
    "5. Score via LightGBM → probability\n"
    "6. Score via LSTM (if temporal data available) → probability\n"
    "7. Combine via ensemble weights → final risk_score\n"
    "8. Compute SHAP values for top 5 drivers\n"
    "9. Classify into risk tier (stable/watch/critical)\n"
    "10. Store result in risk_scores table + Redis cache\n"
    "11. Write to Cassandra for long-term storage\n"
    "12. Return JSON response in < 100ms"
)

doc.add_page_break()

# ════════════════════════════════════════
# 12. NOTIFICATION DISPATCHER
# ════════════════════════════════════════
doc.add_heading("12. Notification Dispatcher", level=1)
doc.add_paragraph(
    "File: intervention/notification_dispatcher.py\n\n"
    "Production-grade multi-channel notification system that ACTUALLY sends "
    "notifications (or runs in simulation mode when credentials aren't configured)."
)

doc.add_heading("12.1. Channel Routing Matrix", level=2)
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Risk Tier"
hdr[1].text = "Score Range"
hdr[2].text = "Channels"
hdr[3].text = "Action"
hdr[4].text = "Urgency"
routing_data = [
    ("Low Watch", "0.30–0.50", "Push, SMS", "Budget nudge", "Low"),
    ("Medium Watch", "0.50–0.65", "SMS, WhatsApp, Push", "Wellness tips", "Medium"),
    ("High Watch", "0.65–0.75", "Email, SMS, WhatsApp", "EMI restructuring", "High"),
    ("Critical", "0.75–0.85", "RM Call, Email, SMS", "Proactive outreach", "Critical"),
    ("Severe", "0.85–1.00", "Collector, RM, Email", "Collections escalation", "Severe"),
]
for tier, rng, channels, action, urgency in routing_data:
    row = table.add_row().cells
    row[0].text = tier
    row[1].text = rng
    row[2].text = channels
    row[3].text = action
    row[4].text = urgency

doc.add_heading("12.2. Email (SMTP)", level=2)
doc.add_paragraph(
    "Uses Python's smtplib to send branded HTML emails via Gmail SMTP or SendGrid.\n"
    "• Branded Barclays template with dark blue header\n"
    "• CTA button linking to wellness page\n"
    "• Professional footer with disclaimers\n"
    "• TLS encryption for all connections"
)

doc.add_heading("12.3. SMS & WhatsApp (Twilio)", level=2)
doc.add_paragraph(
    "SMS: Twilio REST API sends 160-char messages to Indian phone numbers (+91).\n"
    "WhatsApp: Twilio WhatsApp Business API sends messages with emojis and CTAs.\n"
    "• Message SID tracked for delivery confirmation\n"
    "• Cooldown enforcement: no repeat messages within 24 hours per channel"
)

doc.add_heading("12.4. RM Call Assignment", level=2)
doc.add_paragraph(
    "Creates a task in the rm_tasks PostgreSQL table with:\n"
    "• Priority: P1 (score ≥ 0.8) or P2\n"
    "• Due by: NOW + 4 hours\n"
    "• Call script: GenAI-generated with talking points, objection handling\n"
    "• SHAP drivers included for RM context\n"
    "• Optional CRM webhook integration"
)

doc.add_heading("12.5. Collector Escalation", level=2)
doc.add_paragraph(
    "For severe cases (score ≥ 0.85), creates collector assignment with:\n"
    "• Auto-calculated restructuring offer based on DTI and risk score:\n"
    "  - EMI reduction: up to 50%\n"
    "  - Tenure extension: 6–12 months\n"
    "  - Payment holiday: 1–3 months\n"
    "  - Interest rate concession: 25–50 bps\n"
    "  - Settlement offer: 70–90% of outstanding\n"
    "• Full collector brief with financial profile, life events, stress triggers\n"
    "• Optional collections system webhook integration"
)

doc.add_page_break()

# ════════════════════════════════════════
# 13. GENAI
# ════════════════════════════════════════
doc.add_heading("13. GenAI Message Generation (Groq LLM)", level=1)
doc.add_paragraph(
    "File: intervention/genai_messages.py\n\n"
    "Uses Groq's Llama 3.3 70B model to generate personalized messages for each channel. "
    "Each channel has a tailored system prompt:\n\n"
    "• SMS prompt: Under 160 chars, warm, CTA included, Indian English\n"
    "• Email prompt: 3–5 sentences, professional, include specific offer details\n"
    "• WhatsApp prompt: 2–3 lines, friendly, max 2 emojis\n"
    "• RM Script prompt: Structured brief with opening, talking points, objection handling\n"
    "• Collector Brief prompt: Factual, structured, empathetic pre-delinquency tone\n\n"
    "All prompts enforce Barclays brand voice and NEVER mention 'risk', 'delinquency', "
    "or 'default' to the customer. Fallback templates exist for when Groq is unavailable."
)

# ════════════════════════════════════════
# 14. RULES ENGINE
# ════════════════════════════════════════
doc.add_heading("14. Rules Engine & Intervention Logic", level=1)
doc.add_paragraph(
    "File: intervention/rules_engine.py\n\n"
    "The rules engine uses SHAP top drivers to determine the intervention TYPE "
    "(not just the channel). This is what makes the system intelligent:\n\n"
    "SHAP Feature → Intervention Mapping:\n"
    "• salary_delay_days → Payment holiday offer\n"
    "• savings_balance_pct_change_7d → EMI restructuring\n"
    "• lending_app_txn_count_7d → Wellness check-in\n"
    "• failed_autodebits_count_7d → EMI restructuring\n"
    "• discretionary_spend_7d → Budget nudge\n"
    "• gambling_lottery_spend_7d → Wellness check-in\n\n"
    "Interventions are only triggered on risk tier TRANSITIONS (not every score). "
    "Redis tracks the previous tier per customer. This prevents alert fatigue."
)

doc.add_page_break()

# ════════════════════════════════════════
# 15-18. SUPPORTING SYSTEMS
# ════════════════════════════════════════
doc.add_heading("15. Dashboard (Plotly Dash)", level=1)
doc.add_paragraph(
    "File: dashboard/app.py\n\n"
    "Real-time monitoring dashboard at http://localhost:8050:\n"
    "• Risk distribution pie chart (stable/watch/critical)\n"
    "• Score histogram across all customers\n"
    "• Top risk drivers bar chart\n"
    "• Recent interventions table\n"
    "• Auto-refresh every 30 seconds"
)

doc.add_heading("16. Workflow Orchestration", level=1)
doc.add_paragraph(
    "n8n (http://localhost:5678):\n"
    "• Visual workflow for demo: Schedule → Generate Transaction → Score → Risk Check → Alert\n"
    "• Importable workflow: n8n/pdi_demo_workflow.json\n\n"
    "Airflow (http://localhost:8085):\n"
    "• Production DAGs for scheduled batch processing\n"
    "• Model retraining scheduler\n"
    "• Data quality checks"
)

doc.add_heading("17. Monitoring (Grafana + Prometheus)", level=1)
doc.add_paragraph(
    "Grafana (http://localhost:3000):\n"
    "• Scoring service latency (p50, p95, p99)\n"
    "• Request throughput\n"
    "• Error rate monitoring\n\n"
    "Prometheus (http://localhost:9090):\n"
    "• Auto-instrumented FastAPI metrics\n"
    "• Custom metrics: scores_computed, interventions_sent"
)

doc.add_heading("18. CDC Pipeline (Debezium)", level=1)
doc.add_paragraph(
    "Debezium Connect captures PostgreSQL WAL changes in real-time:\n"
    "• Tables monitored: customers, transactions, risk_scores\n"
    "• Changes published to Kafka topics prefixed with 'pdi'\n"
    "• Uses pgoutput plugin and replication slot 'pdi_slot'\n"
    "• Enables real-time downstream analytics in ClickHouse/Superset"
)

doc.add_page_break()

# ════════════════════════════════════════
# 19. DEPLOYMENT
# ════════════════════════════════════════
doc.add_heading("19. Deployment Architecture (Docker + K8s)", level=1)
doc.add_paragraph("16 Docker containers orchestrated via docker-compose.yml:")
containers = [
    ("pdi-postgres", "PostgreSQL 16", "5432"),
    ("pdi-kafka", "Confluent Kafka 7.6", "29092"),
    ("pdi-zookeeper", "Zookeeper", "2181"),
    ("pdi-redis", "Redis 7", "6379"),
    ("pdi-flink-jobmanager", "Apache Flink 1.18", "8081"),
    ("pdi-flink-taskmanager", "Flink TaskManager", "—"),
    ("pdi-spark-master", "Apache Spark 3.5", "8082"),
    ("pdi-spark-worker", "Spark Worker", "—"),
    ("pdi-debezium", "Debezium Connect 2.5", "8083"),
    ("pdi-mlflow", "MLflow 2.11", "5000"),
    ("pdi-grafana", "Grafana 10.3", "3000"),
    ("pdi-prometheus", "Prometheus 2.49", "9090"),
    ("pdi-superset", "Apache Superset 3.1", "8088"),
    ("pdi-airflow", "Apache Airflow 2.9", "8085"),
    ("pdi-cassandra", "Cassandra 4.1", "9042"),
    ("pdi-n8n", "n8n (Workflow Automation)", "5678"),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Container"
hdr[1].text = "Service"
hdr[2].text = "Port"
for name, svc, port in containers:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = svc
    row[2].text = port

doc.add_page_break()

# ════════════════════════════════════════
# 20. TECH STACK
# ════════════════════════════════════════
doc.add_heading("20. Technology Stack Summary", level=1)
categories = [
    ("Languages", "Python 3.10"),
    ("ML/DL", "XGBoost, LightGBM, PyTorch (LSTM), scikit-learn"),
    ("Explainability", "SHAP, LIME"),
    ("Fairness", "Fairlearn, AIF360"),
    ("GenAI", "Groq Cloud (Llama 3.3 70B)"),
    ("Streaming", "Apache Kafka, Apache Flink"),
    ("Batch", "Apache Spark"),
    ("Feature Store", "Feast"),
    ("Database", "PostgreSQL 16, Cassandra 4.1, ClickHouse 23"),
    ("Cache", "Redis 7"),
    ("API", "FastAPI + Uvicorn"),
    ("Dashboard", "Plotly Dash"),
    ("Notifications", "SMTP (Gmail), Twilio (SMS/WhatsApp)"),
    ("Orchestration", "n8n, Apache Airflow"),
    ("Experiment Tracking", "MLflow"),
    ("Monitoring", "Grafana, Prometheus"),
    ("CDC", "Debezium"),
    ("Containerization", "Docker, Docker Compose, Kubernetes"),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Category"
hdr[1].text = "Technology"
for cat, tech in categories:
    row = table.add_row().cells
    row[0].text = cat
    row[1].text = tech

# ════════════════════════════════════════
# 21. DATA FLOW
# ════════════════════════════════════════
doc.add_heading("21. Data Flow Diagram", level=1)
doc.add_paragraph(
    "End-to-End Data Flow:\n\n"
    "┌─────────────┐    ┌───────────┐    ┌──────────────┐    ┌───────────────┐\n"
    "│  Customer    │───→│ PostgreSQL│───→│   Debezium   │───→│    Kafka      │\n"
    "│  Generator   │    │  (pdi_db) │    │   (CDC)      │    │  (Topics)     │\n"
    "└─────────────┘    └───────────┘    └──────────────┘    └───────┬───────┘\n"
    "                                                                │\n"
    "                    ┌─────────────────────────────────────────────┤\n"
    "                    │                                             │\n"
    "                    ▼                                             ▼\n"
    "            ┌──────────────┐                            ┌──────────────┐\n"
    "            │ Flink Stream │                            │ Spark Batch  │\n"
    "            │ (7d/30d      │                            │ (3m features)│\n"
    "            │  features)   │                            │              │\n"
    "            └──────┬───────┘                            └──────┬───────┘\n"
    "                   │                                           │\n"
    "                   └──────────────┬────────────────────────────┘\n"
    "                                  ▼\n"
    "                          ┌──────────────┐\n"
    "                          │  Feast       │\n"
    "                          │ Feature Store│\n"
    "                          └──────┬───────┘\n"
    "                                 ▼\n"
    "                     ┌───────────────────────┐\n"
    "                     │   FastAPI Scoring     │\n"
    "                     │ XGBoost + LightGBM    │\n"
    "                     │ + LSTM Ensemble       │\n"
    "                     │ + SHAP + LIME         │\n"
    "                     └───────────┬───────────┘\n"
    "                                 ▼\n"
    "                     ┌───────────────────────┐\n"
    "                     │   Rules Engine        │\n"
    "                     │ (SHAP → Intervention) │\n"
    "                     └───────────┬───────────┘\n"
    "                                 ▼\n"
    "                     ┌───────────────────────┐\n"
    "                     │ Notification          │\n"
    "                     │ Dispatcher             │\n"
    "                     │ SMS│Email│WA│RM│Coll  │\n"
    "                     └───────────────────────┘"
)

doc.add_page_break()

# ════════════════════════════════════════
# 22. API REFERENCE
# ════════════════════════════════════════
doc.add_heading("22. API Reference", level=1)
doc.add_heading("POST /score", level=2)
doc.add_paragraph(
    "Request:\n"
    '{\n  "customer_id": "CUST_ABC123DEF456"\n}\n\n'
    "Response:\n"
    '{\n'
    '  "customer_id": "CUST_ABC123DEF456",\n'
    '  "risk_score": 0.7234,\n'
    '  "risk_tier": "critical",\n'
    '  "xgboost_score": 0.7456,\n'
    '  "lightgbm_score": 0.7123,\n'
    '  "lstm_score": 0.6890,\n'
    '  "ensemble_score": 0.7234,\n'
    '  "shap_drivers": [\n'
    '    {"feature": "salary_delay_days", "value": 12, "shap_value": 0.183},\n'
    '    {"feature": "lending_app_txn_count_7d", "value": 5, "shap_value": 0.142}\n'
    '  ],\n'
    '  "scored_at": "2026-03-25T01:00:00"\n'
    '}'
)

doc.add_heading("GET /health", level=2)
doc.add_paragraph(
    "Response:\n"
    '{\n'
    '  "status": "healthy",\n'
    '  "models_loaded": {\n'
    '    "xgboost": true,\n'
    '    "lightgbm": true,\n'
    '    "lstm": true,\n'
    '    "shap": true\n'
    '  }\n'
    '}'
)

# ════════════════════════════════════════
# SAVE
# ════════════════════════════════════════
output_path = os.path.join(
    r"c:\Users\L Melvin Denish\barclays\PreDelinquencyEngine",
    "SYSTEM_ARCHITECTURE.docx"
)
doc.save(output_path)
print(f"✅ DOCX saved to: {output_path}")
